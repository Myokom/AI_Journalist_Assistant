from helper_functions import read_docx, read_pdf, handle_other_option, generate_article, generate_title, generate_audio

import streamlit as st
import openai
from docx import Document
from gtts import gTTS

import io
import os
import time
import pdfplumber
import docx2txt


def main():

    st.set_page_config(
        page_title="AI Journalist Assistant",
        page_icon="📰",
        layout="wide",
        initial_sidebar_state="expanded",
    )

    hide_decoration_bar_style = '''
    <style>
        header {visibility: hidden;}
    </style>
    '''
    st.markdown(hide_decoration_bar_style, unsafe_allow_html=True)

    st.title("🖋️ AI Journalist Assistant")
    st.subheader("A tool to help you draft your article and generate a catchy title for your next big story!")

    if 'article' not in st.session_state:
        st.session_state.article = ""
    if 'title' not in st.session_state:
        st.session_state.title = ""


    with st.sidebar:
        # Text box for direct input
        st.subheader("Enter your notes here:")
        text_input = st.text_area(label="User input", label_visibility="collapsed", height=75)

        # File upload option
        st.subheader("... and/or upload a file:")
        uploaded_file = st.file_uploader(label="Upload file", label_visibility="collapsed", type=['txt', 'docx', 'pdf'])

        file_text = ""
        # File processing
        if uploaded_file is not None:
            if uploaded_file.type == "text/plain":
                # Read text file
                file_text = str(uploaded_file.read(), "utf-8")
            elif uploaded_file.type == "application/pdf":
                # Read PDF file
                file_text = read_pdf(uploaded_file)
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                # Read Word file
                file_text = read_docx(uploaded_file)
            else:
                st.error("Unsupported file type")

            st.subheader("File Contents:")
            with st.expander("Show/Hide"):
                st.write(file_text)

        # Combine text input and file text
        user_input = text_input + "\n" + file_text

        st.header("Article Details", divider='red')
        st.subheader("Enter the details of your title:")


        title_style = st.selectbox("Choose the Style for the Article Title", 
                            ["Clickbait", "Simple", "Detailed", "Question-Based", 
                                "Statement", "Humorous", "Shocking", "Inspirational",
                                "How-To", "Listicle", "Other"])
        title_style = handle_other_option(title_style, 'title_style')

        title_lenght = st.slider("Choose the approx. amount of words the Title should be:", 1, 20, 5)

        st.subheader("Enter the details of your article:")


        # Article Type
        article_type = st.selectbox("Select Article Type", 
                                            ["News Report", "Feature Story", "Opinion Piece", "Interview",
                                            "Investigative Article", "Review", "How-to Article", "Profile Piece",
                                            "Editorial", "Column", "Other"])
        article_type = handle_other_option(article_type, 'article_type')

        # Tone and Style
        tone_style = st.selectbox("Select Tone and Style", 
                                    ["Formal", "Informal", "Persuasive", "Informative", "Descriptive",
                                    "Narrative", "Academic", "Journalistic", "Conversational", "Satirical", "Other"])
        tone_style = handle_other_option(tone_style, 'tone_style')

        # Target Audience
        target_audience = st.selectbox("Select Target Audience", 
                                        ["General Public", "Professionals/Business", "Academics/Researchers",
                                            "Teenagers/Young Adults", "Policy Makers", "Specific Interest Groups", "Other"])
        target_audience = handle_other_option(target_audience, 'target_audience')


    with st.form(key='api_submite_form'):

        openai_api_key = st.text_input("Enter your OpenAI API Key:", placeholder="sk-XXXXXXXXXX", type="password")

        submit_button = st.form_submit_button(label='Generate Article')

    if submit_button:
        if openai_api_key:
            with st.status("Working magic...", expanded=True) as status:
                st.write("Sniffing notes... 👃")
                time.sleep(2)
                st.write("Cooking title... 🍳")
                time.sleep(1)
                st.write("Brewing article... 🍵")

                # Generate article and title
                st.session_state.article  = generate_article(user_input, article_type, tone_style, target_audience, openai_api_key)
                st.session_state.title = generate_title(st.session_state.article, title_style, title_lenght, openai_api_key)

                
                # Update status
                status.update(label="Ta-da! All done! 🌟", state="complete", expanded=False)
        else: st.warning("Please enter your OpenAI API Key to generate articles.")
        
    st.subheader(st.session_state.title)
            #st.write_stream(stream_data(article)) 
    st.markdown(st.session_state.article)

    if st.session_state.article and st.session_state.title: 
        st.divider()
        st.subheader("Article Options")   
            
    col1, col2 = st.columns(2)
    with col1:
        # Check if the article and title are available in session state
        if st.session_state.article and st.session_state.title:
            st.subheader("Convert Article to Word Document")
            # Create a Word document
            doc = Document()
            doc.add_heading(st.session_state.title, level=1)
            doc.add_paragraph(st.session_state.article)

            # Save the document to a BytesIO object
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)

            # Generate a file name using the title
            safe_title = "".join([c for c in st.session_state.title if c.isalpha() or c.isdigit() or c==' ']).rstrip()
            file_name = f"{safe_title}.docx"

            # Create a download button
            st.download_button(label="Download Article as Word File",
                            data=file_stream,
                            file_name=file_name,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    with col2:
        if st.session_state.article and st.session_state.title:
            st.subheader("Convert Article to Audio")
            if st.button("Generate Audio"):
                with st.spinner('Generating audio, please wait...'):
                    audio_file_path = generate_audio(st.session_state.article)
                    st.success("Audio file generated successfully!")
                    audio_file = open(audio_file_path, 'rb')
                    st.audio(audio_file.read(), format='audio/mpeg')

                    # Clean up: remove the temp audio file after use
                    os.remove(audio_file_path)


        
            

        #    # Generate and display image based on title
        #    with st.spinner("Generating image based on title..."):
        #                image_url = generate_image(title)
        #                st.image(image_url, caption="Image based on the article title")
        #                st.balloons()




if __name__ == "__main__":
    main()
